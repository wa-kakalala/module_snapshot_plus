from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches
def insert_mdpic(doc,pic_path,pic_name = "xxx"):
    # 添加一个段落，用于插入图片
    paragraph = doc.add_paragraph()
 
    # 插入图片，指定图片路径和宽度（可选）
    run = paragraph.add_run()
    run.add_picture(pic_path,width=Inches(6.0))  # 宽度设置为2英寸
 
    # 获取图片所在的段落，并设置对齐方式为居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 添加一个 run 用于存放题注内容
    caption_para = doc.add_paragraph()
    caption_run = caption_para.add_run()
    caption_run.font.size = Pt(9)
    caption_run.text = f"图x {pic_name}模块"
    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
