# coding: utf-8 
# author: yyrwkk 

from docx.oxml import OxmlElement  
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
gray_color = "bfbfbf"

def create_table(doc,tb_data,tb_header = ["信号名","位宽(bit)","I/O","描述"],tb_name = "xxx"):
    # 添加一个段落用于存放表格题注
    caption_para = doc.add_paragraph()
    
    table = doc.add_table(len(tb_data)+1,len(tb_header), style="Table Grid")

    # 添加一个 run 用于存放题注内容
    caption_run = caption_para.add_run()
    caption_run.font.size = Pt(9)
    caption_run.text = f"表x {tb_name}模块端口表"
    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 将该 run 的样式设置为题注样式（需要自行定义样式或者使用默认样式）
    # caption_run.style = doc.styles["Caption"]
    for col in range( len( tb_header)):
        table.cell(0, col).text = tb_header[col]
        # table.cell(0, col).shading.fill = gray_color
        # 获取单元格的XML元素  
        tc = table.cell(0, col)._tc  
        tcPr = tc.get_or_add_tcPr() 
        # 设置单元格背景颜色为蓝色  
        shd = OxmlElement('w:shd')  
        shd.set(qn('w:val'), 'pct100')  # 设置填充类型为100%填充  
        shd.set(qn('w:fill'), gray_color)  # 设置填充颜色为蓝色（RGB）  
                    
        # 移除现有的shd元素（如果有的话）  
        existing_shds = tcPr.xpath('.//w:shd')  
        for existing_shd in existing_shds:  
            tcPr.remove(existing_shd)  
                    
        # 添加新的背景色元素到单元格属性中  
        tcPr.append(shd)

        table.cell(0, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
        table.cell(0, col).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER # 水平居中
    # 添加表格内容
    for row in range(len( tb_data )):
        for col in range(len(tb_data[row])):
            tb_cell = table.cell(row+1, col)
            tb_cell.text = tb_data[row][col]
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中


    return table